from __future__ import annotations

from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUT = ROOT / "Session7_Kronecker_Embeddings_Guide.docx"
ASSET_DIR = ROOT / "guide_assets"
ASSET_DIR.mkdir(exist_ok=True)


# Compact reference guide preset, resolved to exact values.
PAGE_W = Inches(8.5)
PAGE_H = Inches(11)
MARGIN = Inches(1)
CONTENT_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGIN_DXA = {"top": 80, "bottom": 80, "start": 120, "end": 120}

FONT = "Arial"
MONO = "Courier New"
NAVY = "16324F"
BLUE = "21618C"
TEAL = "167D88"
INK = "1F2933"
MUTED = "5C6B73"
LIGHT_BLUE = "EAF3F8"
LIGHT_TEAL = "E8F5F5"
LIGHT_GOLD = "FFF7DF"
LIGHT_GRAY = "F3F5F7"
MID_GRAY = "D6DDE2"
WHITE = "FFFFFF"
RED = "A23B3B"
GREEN = "2D6A4F"


def rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value)


def set_cell_margins(cell, **kwargs):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side in ("top", "start", "bottom", "end"):
        if side in kwargs:
            node = tc_mar.find(qn(f"w:{side}"))
            if node is None:
                node = OxmlElement(f"w:{side}")
                tc_mar.append(node)
            node.set(qn("w:w"), str(kwargs[side]))
            node.set(qn("w:type"), "dxa")


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = MID_GRAY, size: int = 6):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        elem = borders.find(tag)
        if elem is None:
            elem = OxmlElement(f"w:{edge}")
            borders.append(elem)
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:sz"), str(size))
        elem.set(qn("w:color"), color)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_table_geometry(table, widths_dxa: list[int], indent_dxa: int = TABLE_INDENT_DXA):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        prevent_row_split(row)
        for index, cell in enumerate(row.cells):
            width = widths_dxa[min(index, len(widths_dxa) - 1)]
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell, **CELL_MARGIN_DXA)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_border(cell)


def paragraph_shading(paragraph, fill: str):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def paragraph_left_border(paragraph, color: str = TEAL, size: int = 18, space: int = 8):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(size))
    left.set(qn("w:space"), str(space))
    left.set(qn("w:color"), color)
    p_bdr.append(left)


def set_run(run, *, font=FONT, size=11, color=INK, bold=False, italic=False):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.font.color.rgb = rgb(color)
    run.bold = bold
    run.italic = italic
    return run


def add_field(paragraph, field_code: str):
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = field_code
    fld_char_sep = OxmlElement("w:fldChar")
    fld_char_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char_begin, instr_text, fld_char_sep, text, fld_char_end])
    set_run(run, size=9, color=MUTED)


def configure_styles(doc: Document):
    styles = doc.styles

    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.18
    normal.paragraph_format.widow_control = True

    title = styles["Title"]
    title.font.name = FONT
    title._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    title._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    title.font.size = Pt(30)
    title.font.bold = True
    title.font.color.rgb = rgb(NAVY)
    title.paragraph_format.space_after = Pt(8)

    subtitle = styles["Subtitle"]
    subtitle.font.name = FONT
    subtitle._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    subtitle._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    subtitle.font.size = Pt(14)
    subtitle.font.color.rgb = rgb(BLUE)
    subtitle.paragraph_format.space_after = Pt(10)

    heading_specs = {
        "Heading 1": (17, BLUE, 18, 9),
        "Heading 2": (13.5, BLUE, 13, 6),
        "Heading 3": (11.5, NAVY, 9, 4),
    }
    for name, (size, color, before, after) in heading_specs.items():
        style = styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style.font.size = Pt(10.5)
        style.font.color.rgb = rgb(INK)
        style.paragraph_format.left_indent = Inches(0.38)
        style.paragraph_format.first_line_indent = Inches(-0.19)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.18

    code = styles.add_style("Code Block", 1)
    code.font.name = MONO
    code._element.rPr.rFonts.set(qn("w:ascii"), MONO)
    code._element.rPr.rFonts.set(qn("w:hAnsi"), MONO)
    code.font.size = Pt(8.4)
    code.font.color.rgb = rgb(INK)
    code.paragraph_format.left_indent = Inches(0.16)
    code.paragraph_format.right_indent = Inches(0.16)
    code.paragraph_format.space_before = Pt(4)
    code.paragraph_format.space_after = Pt(8)
    code.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    caption = styles["Caption"]
    caption.font.name = FONT
    caption._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    caption._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = rgb(MUTED)
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(8)


def add_body(doc, text: str, *, after=6, bold_prefix: str | None = None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    if bold_prefix and text.startswith(bold_prefix):
        set_run(p.add_run(bold_prefix), bold=True)
        set_run(p.add_run(text[len(bold_prefix):]))
    else:
        set_run(p.add_run(text))
    return p


def add_bullet(doc, text: str, *, level=0, bold_prefix: str | None = None):
    style = "List Bullet" if level == 0 else "List Bullet 2"
    p = doc.add_paragraph(style=style)
    p.paragraph_format.left_indent = Inches(0.38 + level * 0.24)
    p.paragraph_format.first_line_indent = Inches(-0.19)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.18
    if bold_prefix and text.startswith(bold_prefix):
        set_run(p.add_run(bold_prefix), bold=True)
        set_run(p.add_run(text[len(bold_prefix):]))
    else:
        set_run(p.add_run(text))
    return p


def add_number(doc, text: str, *, bold_prefix: str | None = None):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.38)
    p.paragraph_format.first_line_indent = Inches(-0.19)
    p.paragraph_format.space_after = Pt(4)
    if bold_prefix and text.startswith(bold_prefix):
        set_run(p.add_run(bold_prefix), bold=True)
        set_run(p.add_run(text[len(bold_prefix):]))
    else:
        set_run(p.add_run(text))
    return p


def add_callout(doc, label: str, text: str, *, fill=LIGHT_TEAL, border=TEAL):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.right_indent = Inches(0.12)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15
    paragraph_shading(p, fill)
    paragraph_left_border(p, border)
    set_run(p.add_run(label.upper() + "  "), size=9.3, color=border, bold=True)
    set_run(p.add_run(text), size=10.2, color=INK)
    return p


def add_code(doc, code_text: str):
    p = doc.add_paragraph(style="Code Block")
    p.paragraph_format.keep_together = True
    paragraph_shading(p, LIGHT_GRAY)
    for index, line in enumerate(code_text.strip("\n").splitlines()):
        if index:
            p.add_run().add_break()
        set_run(p.add_run(line), font=MONO, size=8.4, color=INK)
    return p


def add_table(doc, headers: list[str], rows: list[list[str]], widths: list[int], *, first_col_bold=False):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, text in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_shading(cell, NAVY)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_run(p.add_run(text), size=9.3, color=WHITE, bold=True)
    for row_data in rows:
        row = table.add_row()
        for i, text in enumerate(row_data):
            cell = row.cells[i]
            if len(table.rows) % 2 == 0:
                set_cell_shading(cell, "F8FAFB")
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.08
            set_run(p.add_run(text), size=9.2, color=INK, bold=(first_col_bold and i == 0))
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_section_break(doc):
    doc.add_section(WD_SECTION.NEW_PAGE)


def add_source_note(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(0)
    set_run(p.add_run("Source basis: "), size=8.5, color=MUTED, bold=True)
    set_run(
        p.add_run("Session7_Transcript.md and assignment.md in the session7 folder. The transcript is machine-generated and contains terminology/number transcription errors."),
        size=8.5,
        color=MUTED,
    )


def make_pipeline_diagram(path: Path):
    w, h = 1500, 560
    image = Image.new("RGB", (w, h), "white")
    draw = ImageDraw.Draw(image)
    font_path = "/System/Library/Fonts/Supplemental/Arial.ttf"
    bold_path = "/System/Library/Fonts/Supplemental/Arial Bold.ttf"
    font = ImageFont.truetype(font_path, 31)
    small = ImageFont.truetype(font_path, 25)
    bold = ImageFont.truetype(bold_path, 31)
    title = ImageFont.truetype(bold_path, 34)

    def box(x1, y1, x2, y2, fill, outline, heading, lines):
        draw.rounded_rectangle((x1, y1, x2, y2), 18, fill=fill, outline=outline, width=4)
        draw.text((x1 + 22, y1 + 20), heading, font=bold, fill="#16324F")
        y = y1 + 67
        for line in lines:
            draw.text((x1 + 22, y), line, font=small, fill="#1F2933")
            y += 35

    def arrow(x1, y1, x2, y2, color="#167D88"):
        draw.line((x1, y1, x2, y2), fill=color, width=7)
        draw.polygon([(x2, y2), (x2 - 18, y2 - 13), (x2 - 18, y2 + 13)], fill=color)

    draw.text((55, 25), "Two ways to create the vector that enters the transformer", font=title, fill="#16324F")
    draw.text((55, 100), "Standard learned lookup", font=bold, fill="#21618C")
    box(55, 155, 320, 290, "#EAF3F8", "#21618C", "token id", ["e.g. 47"])
    arrow(330, 222, 435, 222, "#21618C")
    box(445, 155, 790, 290, "#EAF3F8", "#21618C", "V x d table", ["gather row 47"])
    arrow(800, 222, 905, 222, "#21618C")
    box(915, 155, 1425, 290, "#EAF3F8", "#21618C", "d-dimensional vector", ["learned per vocabulary row"])

    draw.text((55, 335), "Kronecker-style construction", font=bold, fill="#167D88")
    box(55, 390, 320, 525, "#E8F5F5", "#167D88", "token text", ["e.g. apple"])
    arrow(330, 457, 435, 457)
    box(445, 390, 790, 525, "#E8F5F5", "#167D88", "byte + position", ["fixed, deterministic"])
    arrow(800, 457, 905, 457)
    box(915, 390, 1425, 525, "#E8F5F5", "#167D88", "d-dimensional vector", ["small learned projection"])
    image.save(path, quality=95)


def set_picture_alt(inline_shape, title: str, description: str):
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("title", title)
    doc_pr.set("descr", description)


def build_document():
    pipeline_path = ASSET_DIR / "embedding_pipeline.png"
    make_pipeline_diagram(pipeline_path)

    doc = Document()
    section = doc.sections[0]
    section.page_width = PAGE_W
    section.page_height = PAGE_H
    section.top_margin = MARGIN
    section.bottom_margin = MARGIN
    section.left_margin = MARGIN
    section.right_margin = MARGIN
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True
    configure_styles(doc)

    # Running header and footer.
    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_run(hp.add_run("SESSION 7  •  EMBEDDINGS & KRONECKER V2"), size=8.5, color=MUTED, bold=True)
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run(fp.add_run("ERA V5  |  "), size=8.5, color=MUTED)
    add_field(fp, "PAGE")

    # Cover.
    for _ in range(4):
        doc.add_paragraph()
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(16)
    set_run(kicker.add_run("SESSION 7 CONCEPT GUIDE"), size=11, color=TEAL, bold=True)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("From Embedding Tables to\nKronecker Embeddings")

    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("The problem, the proposed idea, what broke, and what Kronecker V2 is trying to solve")

    line = doc.add_paragraph()
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    line.paragraph_format.space_before = Pt(14)
    line.paragraph_format.space_after = Pt(22)
    set_run(line.add_run("A plain-language walkthrough with examples and code"), size=10.5, color=MUTED, italic=True)

    cover_callout = doc.add_paragraph()
    cover_callout.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover_callout.paragraph_format.left_indent = Inches(0.55)
    cover_callout.paragraph_format.right_indent = Inches(0.55)
    cover_callout.paragraph_format.space_before = Pt(18)
    cover_callout.paragraph_format.space_after = Pt(40)
    paragraph_shading(cover_callout, LIGHT_TEAL)
    set_run(cover_callout.add_run("Core idea: "), size=11, color=TEAL, bold=True)
    set_run(
        cover_callout.add_run("replace a huge learned row-per-token table with a deterministic construction from the token's own bytes and positions, followed by a much smaller learned projection."),
        size=11,
        color=INK,
    )

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(meta.add_run("Based on the Session 7 transcript and assignment  •  8 August 2026"), size=9.2, color=MUTED)
    doc.add_page_break()

    # Contents / reading path.
    doc.add_heading("How to use this guide", level=1)
    add_body(doc, "If the session felt dense, read Sections 1, 4, 6, and 8 first. They give the story without requiring the supporting mechanics. Return to the other sections when a term such as fertility, weight tying, or positional encoding appears.")
    add_table(
        doc,
        ["Section", "What you will understand"],
        [
            ["1. Session in one page", "The complete problem → idea → limitation → next-step story."],
            ["2–3. Foundations", "How ordinary embeddings work and why their cost matters."],
            ["4–5. Kronecker V1", "The construction, examples, intended benefits, and impact."],
            ["6. What went wrong", "The 32-byte cap, Indic disadvantage, training shock, missing proof, and reverse-decoding problem."],
            ["7. Position information", "Why order must be injected and where learned, sinusoidal, RoPE, and ALiBi differ."],
            ["8–9. Kronecker V2", "The five separate assignment paths and what counts as convincing evidence."],
        ],
        [2050, 7310],
        first_col_bold=True,
    )
    add_callout(doc, "Notation note", "The transcript repeatedly says “8096,” but its own construction uses 32 byte positions × 256 byte values = 8192, and later mentions an 8192 → 4096 projection. This guide uses 8192 as the intended value and treats 8096 as a transcription error.", fill=LIGHT_GOLD, border="9A7200")
    add_source_note(doc)
    doc.add_page_break()

    # Section 1.
    doc.add_heading("1. The session in one page", level=1)
    doc.add_heading("The problem", level=2)
    add_body(doc, "A transformer cannot consume a token ID such as 47 directly as meaning. It first needs a dense vector. The usual solution is a learned embedding table with one row for every vocabulary token. A vocabulary of 131,072 tokens and width 8,192 needs more than one billion learned values before the transformer stack even begins.")
    add_body(doc, "That design creates three pressures at once: the table is expensive to train, rare tokens receive few updates because language is Zipf-distributed, and increasing vocabulary to reduce token count makes the table even larger.")

    doc.add_heading("The proposed idea: Kronecker embedding V1", level=2)
    add_body(doc, "Instead of storing one learned row for every token, construct a fixed vector from the token's UTF-8 bytes and their positions. “apple” is therefore represented using the bytes for a, p, p, l, e together with positions 0, 1, 2, 3, 4. A shared learned projection then adapts this deterministic representation into the vector space the transformer wants.")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shape = p.add_run().add_picture(str(pipeline_path), width=Inches(6.45))
    set_picture_alt(shape, "Standard versus Kronecker embedding pipeline", "The standard path gathers a row from a vocabulary table. The Kronecker path constructs a fixed byte-and-position code from token text and applies a shared learned projection.")
    cap = doc.add_paragraph("Figure 1. The transformer receives a d-dimensional vector in both designs; only the way that vector is created changes.", style="Caption")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("Why it is attractive", level=2)
    for text in [
        "The trainable input embedding no longer grows linearly with vocabulary size.",
        "The model sees token-internal spelling/byte structure instead of an arbitrary token ID alone.",
        "Tokens that share bytes and positions share parts of the construction, so rare or unseen strings may inherit useful structure.",
        "The same token always receives the same base code; context is added later by the transformer.",
    ]:
        add_bullet(doc, text)

    doc.add_heading("What blocked V1", level=2)
    add_body(doc, "The construction used a fixed 32-byte window. Short strings waste capacity; strings longer than 32 UTF-8 bytes are cropped. Indic characters often use three UTF-8 bytes, so the practical character limit can fall to roughly ten. The shared projection also showed loss spikes when the training mixture changed quickly. Most importantly, the forward construction was not reliably reversible: a model predicts an approximate vector, not the exact deterministic code required to decode a token without an output head.")

    add_callout(doc, "What we are solving now", "Kronecker V2 is not one single prescribed design. The assignment asks you to choose one independent limitation—mathematical structure, multimodality, dynamic length, a Fourier construction, or reliable reverse decoding—and propose and experimentally prove a solution.")
    doc.add_page_break()

    # Section 2.
    doc.add_heading("2. Ordinary embeddings: the foundation", level=1)
    doc.add_heading("2.1 A token ID is an address, not meaning", level=2)
    add_body(doc, "A tokenizer converts text into IDs. If “apple” is token 47, the number 47 merely says where to look. The embedding layer turns that address into a vector of d learned numbers.")
    add_code(doc, """# Shape notation
token_ids: [batch, sequence]          # integers
E:         [vocabulary, embedding_d]  # learned table
x = E[token_ids]                      # [batch, sequence, embedding_d]""")
    add_callout(doc, "Operational view", "Textbooks often describe embedding lookup as multiplying a one-hot vector by the whole table. The mathematics is equivalent, but implementations use a gather: select the requested rows. Reading a row is cheap; storing and training the full table is the expensive part.", fill=LIGHT_BLUE, border=BLUE)

    doc.add_heading("2.2 What does a dimension mean?", level=2)
    add_body(doc, "The session used “bandwidth” as the intuition. More dimensions give the system more capacity to represent distinctions and relationships. Individual dimensions usually do not have clean human labels such as “is a company” or “is red.” Neural networks distribute concepts across many coordinates and learn features we would not design by hand.")
    add_body(doc, "For example, the base embedding for “bank” must support contexts involving money, a river bank, or banking a road. The transformer later uses surrounding words to produce the context-specific representation.")

    doc.add_heading("2.3 Forward and backward passes", level=2)
    add_number(doc, "Forward: gather the rows for all token IDs in the batch.", bold_prefix="Forward:")
    add_number(doc, "Backward: accumulate all gradient contributions for each row that appeared.", bold_prefix="Backward:")
    add_number(doc, "Update: the optimizer applies one combined update per touched row for the step.", bold_prefix="Update:")
    add_body(doc, "If “the” appears 67 times, its row receives 67 gradient contributions that are accumulated. A token absent from the batch receives no embedding-row update in that step.")

    doc.add_heading("2.4 Zipf's law creates uneven learning", level=2)
    add_body(doc, "Natural language contains a small number of extremely frequent tokens and a long tail of rare tokens. Therefore frequent function words are touched constantly while domain terms, names, and lower-resource-language tokens may be updated far less often.")
    add_callout(doc, "Impact", "The imbalance is not only between languages; it exists inside every language. Less data for an Indic language plus a long rare-token tail compounds the problem. Optimizers and data mixture design can reduce the harm, but neither makes the long tail disappear.", fill=LIGHT_GOLD, border="9A7200")
    doc.add_page_break()

    # Section 3.
    doc.add_heading("3. The economic and architectural trade-offs", level=1)
    doc.add_heading("3.1 The size of the table", level=2)
    add_body(doc, "For vocabulary size V and embedding dimension d, the input table contains V × d parameters.")
    add_code(doc, """V = 131_072
d = 8_192
parameters = V * d                 # 1,073,741,824
bf16_weights = parameters * 2     # about 2.0 GiB
adamw_training_state ≈ parameters * 16  # about 16 GiB, implementation-dependent""")
    add_body(doc, "An untied output head of the same shape adds another V × d parameters. Training-memory estimates vary with precision, optimizer implementation, sharding, and master-weight policy; “about 16 bytes per trainable parameter” is the session's useful planning heuristic, not a universal constant.")
    add_callout(doc, "Technical clarification", "Parameter memory is not copied once per batch item. Larger batches mainly increase activation and temporary-buffer memory. The embedding weights and optimizer state remain one model copy per replica or shard.", fill=LIGHT_GOLD, border="9A7200")

    doc.add_heading("3.2 Vocabulary size versus fertility", level=2)
    add_body(doc, "Fertility is roughly how many tokens are needed to represent the same text. A larger vocabulary can store longer or more common subwords, reducing fertility and attention compute. But every additional token adds another embedding row and another output-head column.")
    add_table(
        doc,
        ["Choice", "Benefit", "Cost"],
        [
            ["Smaller vocabulary", "Smaller input/output matrices", "More tokens per sentence; more attention work"],
            ["Larger vocabulary", "Fewer tokens; better coverage of words and scripts", "Larger embedding and output matrices"],
        ],
        [1900, 3450, 4010],
        first_col_bold=True,
    )
    add_body(doc, "The correct point depends on the actual language mixture, tokenizer fertility, model width, hardware-friendly sizes, and available memory. Tokenizer and embedding design are therefore one coupled decision surface.")

    doc.add_heading("3.3 Output head and weight tying", level=2)
    add_body(doc, "The transformer produces a d-dimensional hidden state. The output head maps it to V logits, one score per vocabulary token. Weight tying reuses the input embedding matrix (transposed) as this output map.")
    for text in [
        "Tying saves V × d parameters and can regularize small models.",
        "Untying lets input representation and output discrimination specialize independently.",
        "The session's rule of thumb was to tie small language models and prefer untied heads for larger models, based on empirical trade-offs rather than a mathematical necessity.",
    ]:
        add_bullet(doc, text)

    doc.add_heading("3.4 Low-rank factorization", level=2)
    add_body(doc, "A large matrix can sometimes be approximated by two narrower matrices: d × r followed by r × d, where r is the bottleneck rank. This is lossy compression unless the original matrix truly has rank r, but trained networks often contain directions that matter much more than others.")
    add_code(doc, """# Full transform
y = x @ W                 # W: [8192, 8192]

# Low-rank approximation
y = (x @ A) @ B           # A: [8192, r], B: [r, 8192]
# r controls the compute/quality trade-off""")
    add_body(doc, "The session connected this intuition to LoRA. The safe rank is empirical and can change across layers and training stages; it should be measured rather than guessed.")
    doc.add_page_break()

    # Section 4.
    doc.add_heading("4. Kronecker embedding V1, step by step", level=1)
    doc.add_heading("4.1 The motivating question", level=2)
    add_callout(doc, "Question", "Can we produce the same d-dimensional input vector without keeping a V × d row table?", fill=LIGHT_BLUE, border=BLUE)
    add_body(doc, "Kronecker V1 answers: build a structured code directly from the token's bytes and positions, then let one shared trainable projection adapt that code for the model.")

    doc.add_heading("4.2 A clean mental model of the construction", level=2)
    add_body(doc, "UTF-8 provides 256 possible byte values. The design reserves 32 byte positions. Pairing a position with a byte gives 32 × 256 = 8192 possible position-byte features.")
    add_body(doc, "For byte b at position i, take the Kronecker product of a one-hot position vector and a one-hot byte vector. This produces one active coordinate in an 8192-dimensional vector. Sum or average these features across the bytes of the token.")
    add_code(doc, """import torch

def kronecker_code(text: str, max_bytes: int = 32):
    raw = text.encode("utf-8")
    if len(raw) > max_bytes:
        raise ValueError("token exceeds the 32-byte window")

    x = torch.zeros(max_bytes * 256)  # 8192 fixed features
    for position, byte_value in enumerate(raw):
        x[position * 256 + byte_value] = 1.0

    return x / max(1, len(raw))

# Shared learned adapter; no row per vocabulary token
projection = torch.nn.Linear(8192, d_model, bias=False)
embedding = projection(kronecker_code("apple"))""")
    add_body(doc, "This snippet is an explanatory implementation consistent with the session's 32 × 256 construction. The paper/code may implement the fixed mapping with an equivalent structured or random projection, but the conceptual separation remains the same: deterministic code first, shared learned adapter second.")

    doc.add_heading("4.3 Worked example: “apple”", level=2)
    add_table(
        doc,
        ["Byte position", "Character", "UTF-8 byte", "Active feature"],
        [
            ["0", "a", "97", "0 × 256 + 97"],
            ["1", "p", "112", "1 × 256 + 112"],
            ["2", "p", "112", "2 × 256 + 112"],
            ["3", "l", "108", "3 × 256 + 108"],
            ["4", "e", "101", "4 × 256 + 101"],
        ],
        [1500, 1750, 2500, 3610],
    )
    add_body(doc, "The repeated letter p is distinguishable because its position differs. The base code is deterministic: every occurrence of the same byte sequence produces the same vector. The transformer later makes “apple” mean fruit, company, or metaphor based on context.")
    doc.add_page_break()

    # Section 5.
    doc.add_heading("5. What Kronecker V1 is trying to buy us", level=1)
    add_table(
        doc,
        ["Goal", "Why the standard table struggles", "Kronecker promise"],
        [
            ["Lower trainable cost", "Input cost grows as V × d", "Shared projection can be independent of V"],
            ["Expose spelling", "A token ID does not reveal its characters", "Bytes and positions are present in the base code"],
            ["Help rare strings", "Rare rows receive few updates", "Shared byte/position machinery receives learning from many strings"],
            ["Handle novel forms", "Unknown text is usually split into known subwords", "Any byte sequence within the window has a deterministic code"],
            ["Support larger vocabularies", "Input/output matrices become costly", "Input-side growth can be removed; output still needs a solution"],
        ],
        [1850, 3750, 3760],
        first_col_bold=True,
    )
    doc.add_heading("Important boundaries around these claims", level=2)
    add_bullet(doc, "A deterministic encoder alone does not make an arbitrary phrase a token. The tokenizer or a higher-level packing policy must decide which spans are passed as units.")
    add_bullet(doc, "Standard learned embeddings are also deterministic at inference for a given token ID. Kronecker's distinctive property is that the base vector is generated algorithmically from token content rather than stored as a learned vocabulary row.")
    add_bullet(doc, "Character structure may help spelling and related forms, but semantic usefulness must be measured. Similar spelling does not always imply similar meaning.")
    add_bullet(doc, "The session reported encouraging small-scale behavior, around a 121–131 million parameter test, but not conclusive large-scale evidence.")

    doc.add_heading("Why a shared projection is still useful", level=2)
    add_body(doc, "The fixed code answers “what bytes are present where?” It does not automatically arrange concepts in the geometry the transformer prefers. A learned projection can rotate, combine, or ignore fixed features. If the fixed code is already sufficient, the projection can learn something close to an identity-like mapping; otherwise it can adapt it without forcing the first transformer block to do all the work.")
    add_callout(doc, "The design principle used in the session", "Provide a proposed solution as a learnable option, but allow the network to ignore it when it is not useful. This avoids hard-coding the assumption that the engineered feature is always correct.")
    doc.add_page_break()

    # Section 6.
    doc.add_heading("6. What V1 ran into", level=1)
    doc.add_heading("6.1 Fixed 32-byte window", level=2)
    add_body(doc, "Every token is allocated against the same 32-byte design. “a” uses one byte and leaves 31 positions empty; “apple” uses five and leaves 27. A token above 32 bytes must be rejected, split, or cropped. Cropping is especially dangerous because different long strings can become identical after the boundary.")

    doc.add_heading("6.2 UTF-8 is not script-neutral at this boundary", level=2)
    add_body(doc, "ASCII letters typically use one UTF-8 byte. Many Indic characters use three bytes. Therefore the same 32-byte budget may hold about 32 ASCII characters but only about ten Indic characters, sometimes fewer once combining marks are included.")
    add_code(doc, """for s in ["internationalization", "namaste", "नमस्ते"]:
    print(s, len(s), len(s.encode("utf-8")))

# Character count and byte count are different questions.
# The V1 limit is a byte-window limit.""")
    add_callout(doc, "Impact", "A design intended to help Indic languages can accidentally impose a stricter usable word-length limit on them. The tokenizer must either forbid overlong tokens or the embedding must become length-adaptive.", fill=LIGHT_GOLD, border="9A7200")

    doc.add_heading("6.3 Rapid curriculum shifts shock the adapter", level=2)
    add_body(doc, "The session described spikes when training moved abruptly from simpler/common English to more specialized language. The shared projection is a bottleneck touched by every token; a sudden distribution change can produce a sharp gradient shift. The proposed operational mitigation was a gradual mixture transition, warm-up, and suitable learning-rate control.")

    doc.add_heading("6.4 Evidence was not yet conclusive", level=2)
    add_body(doc, "The idea was explored on a small model, but GPU/time limits prevented the larger controlled training needed to establish the claimed advantages. Good end-task predictions are not enough; we also need matched baselines and specific tests for rare tokens, misspellings, multilingual behavior, memory, and throughput.")

    doc.add_heading("6.5 The reverse path is the hardest blocker", level=2)
    add_body(doc, "Forward encoding is exact: token → deterministic code. Generation needs the reverse: model state → token. A trained model rarely emits the exact target code. If the target is [0.30, 0.20, …], the model might produce [0.31, 0.18, …]. Without a classifier/output head, the system must still decide which token that noisy point represents.")
    for text in [
        "Exact equality is unusable with continuous neural outputs.",
        "Nearest-neighbor or cosine search requires a codebook and can be unreliable early in training when outputs are far from all valid codes.",
        "Searching a million candidates can recreate a large-vocabulary cost in a different form.",
        "A probabilistic neighborhood loss (the session mentions KL divergence / VAE-style distributions) can add tolerance, but it does not by itself guarantee unique, efficient decoding.",
    ]:
        add_bullet(doc, text)
    add_callout(doc, "Why this matters", "A reliable reverse codec could remove the V-way output head and make very large or open vocabularies practical. Until then, Kronecker mainly solves the input-side table, not the full language-model interface.", fill=LIGHT_BLUE, border=BLUE)
    doc.add_page_break()

    # Section 7.
    doc.add_heading("7. Position information: a separate but related topic", level=1)
    add_body(doc, "A transformer processes the tokens in a sequence in parallel. Without position information, its core operations cannot distinguish “A before B” from “B before A” purely from the input set. Position must therefore be injected somewhere.")
    add_table(
        doc,
        ["Family", "Where/how", "Main limitation or benefit"],
        [
            ["Learned absolute", "A learned vector per sequence position, added near input", "Positions beyond training are unlearned; table grows with maximum context"],
            ["Sinusoidal", "Fixed sine/cosine functions added near input", "No learned table, but long-range resolution/extrapolation can be weak"],
            ["RoPE", "Rotates query/key features inside attention", "Encodes relative position effectively; modern common choice"],
            ["ALiBi", "Adds distance-dependent bias to attention scores", "Simple and cheap; directly favors nearer positions"],
        ],
        [1700, 3500, 4160],
        first_col_bold=True,
    )
    add_body(doc, "Kronecker V1 also needs token-internal position: the p at byte position 1 must differ from the p at position 2. That local byte position is not the same as the token's position in the sentence. The two levels should not be confused.")
    add_callout(doc, "Two position systems", "Inside-token position tells us where a byte occurs within “apple.” Sequence position tells us where the token occurs within the sentence. A complete model may need both, implemented in different places.")
    doc.add_page_break()

    # Section 8.
    doc.add_heading("8. Kronecker V2: the five independent assignment directions", level=1)
    add_body(doc, "The assignment explicitly says to pick one problem, not combine all five. Each direction asks for a proposed mechanism plus a small-model experiment that could falsify it.")

    doc.add_heading("8.1 Mathematical structure inside embeddings", level=2)
    add_body(doc, "Goal: reserve part of the embedding so mathematical operations have predictable geometric effects—for example, a representation of 9 plus a representation of 9 maps to 18.")
    add_body(doc, "The precise concept is an operation-preserving representation. Addition can be encoded linearly in a value subspace. Multiplication is harder to make linear in the same coordinates; positive multiplication becomes addition after a log transform, or it can be represented by a learned operator conditioned on ×.")
    add_code(doc, """# One simple research sketch, not a complete solution
e(n) = concat(language_features(n), numeric_value(n), log_value(n))

addition_loss       = || value(e(a) + e(b)) - (a + b) ||
multiplication_loss = || logpart(e(a)) + logpart(e(b)) - log(a*b) ||
language_loss       = next_token_cross_entropy(...)""")
    add_body(doc, "Proof should test held-out numbers and operations, not only memorize training pairs. Compare arithmetic accuracy and normal language modeling against a same-size baseline.")

    doc.add_heading("8.2 One construction for text, images, and audio", level=2)
    add_body(doc, "Goal: extend the “content code + local position + shared projection” idea beyond bytes. Text units could be bytes, image units could be patch codes, and audio units could be short-time spectral or codec codes. Add a modality identifier so identical numeric codes from different modalities do not collide.")
    add_body(doc, "A minimal test should include matched classification or retrieval tasks for all three modalities and at least one cross-modal task. The important question is whether the shared construction transfers useful structure without erasing modality-specific detail.")

    doc.add_heading("8.3 Dynamic length instead of a 32-byte box", level=2)
    add_body(doc, "Goal: make cost depend on actual length and remove hard cropping. Candidate directions include sparse gathers over position-byte features, chunked/hierarchical composition, recurrent hashing, or a variable-length pooling encoder.")
    add_body(doc, "A good design must preserve order, avoid collisions, support long strings, and remain computationally bounded. Test English and Indic strings separately, reporting bytes per character, truncation rate, throughput, and downstream quality.")
    doc.add_page_break()

    doc.add_heading("8.4 A real Fourier alternative", level=2)
    add_body(doc, "Goal: represent each character/byte as frequencies and encode position through phase, then combine a word by adding waves. Fourier addition is fast and naturally compositional, but naïve summation can lose order or create collisions.")
    add_code(doc, """# Conceptual complex-valued code
code(token) = Σ_i amplitude(byte_i) * exp(j * frequency(byte_i) * position_i)

# Questions the experiment must answer:
# 1. Can two different strings collide or become too similar?
# 2. Can the original sequence be recovered under prediction noise?
# 3. Does the code generalize beyond lengths seen in training?""")
    add_body(doc, "The strongest proof is not a pretty frequency plot. It is a collision analysis plus reconstruction and language-model performance under controlled noise.")

    doc.add_heading("8.5 Reverse / invertible Kronecker", level=2)
    add_body(doc, "Goal: decode the model's approximate continuous output into the original token reliably enough to remove the V-way output head. This was identified in the session as the most immediately valuable independent problem.")
    for text in [
        "Error-correcting structured codes: valid tokens occupy separated regions with known decoding rules.",
        "Autoregressive byte decoder: predict length and bytes from the model state; compute depends on output length rather than vocabulary size.",
        "Vector quantization: snap outputs to learned or deterministic code regions, with straight-through or commitment losses.",
        "Probabilistic code distributions: predict a mean/variance or categorical factors, then decode within a tolerance region.",
    ]:
        add_bullet(doc, text)
    add_body(doc, "A successful reverse method must work from random initialization through training—not only after a mature model has already learned to land near valid codes. Measure exact token recovery, robustness to noise, decoding latency, memory, and next-token quality.")
    add_callout(doc, "Do not overclaim", "Solving reverse decoding does not automatically make a one-million-token system “free.” Tokenizer construction, decoding search, training stability, and data coverage still matter. It can, however, remove the dominant dense output matrix if the decoder is genuinely sublinear in V.", fill=LIGHT_GOLD, border="9A7200")
    doc.add_page_break()

    # Section 9.
    doc.add_heading("9. What counts as proof for the assignment", level=1)
    add_body(doc, "The assignment asks for code that demonstrates the chosen mechanism. A small transformer is enough if the experiment isolates the claim and uses a fair baseline.")
    doc.add_heading("A minimal experimental recipe", level=2)
    add_number(doc, "State one hypothesis in measurable form. Example: “a dynamic encoder has zero truncation and equal or better validation loss than fixed-32 Kronecker at the same trainable-parameter budget.”")
    add_number(doc, "Implement three systems: ordinary learned embedding baseline, Kronecker V1 baseline, and your proposed V2 change.")
    add_number(doc, "Match the important budget: model width, transformer depth, training tokens, optimizer, steps, and as closely as possible trainable parameters or FLOPs.")
    add_number(doc, "Use a dataset designed to expose the target problem, then keep a held-out split that contains unseen combinations—not just unseen examples copied from the same templates.")
    add_number(doc, "Report quality, efficiency, and failure cases. A single final loss is not enough.")
    add_number(doc, "Repeat with several seeds or report variability. Small neural experiments can look successful by chance.")

    doc.add_heading("Metrics by problem", level=2)
    add_table(
        doc,
        ["Chosen problem", "Primary evidence", "Necessary stress test"],
        [
            ["Math structure", "Held-out operation accuracy; embedding consistency", "Numbers/ranges/compositions absent from training"],
            ["Multimodal", "Per-modality quality and cross-modal retrieval", "Ablate modality and position codes"],
            ["Dynamic length", "No truncation; quality vs length; speed/memory", "Long Indic and ASCII strings; collision tests"],
            ["Fourier", "Reconstruction/collision rate and LM quality", "Noise, long sequences, permutations/anagrams"],
            ["Reverse decoding", "Exact recovery and next-token accuracy", "Early training, output noise, very large candidate set"],
        ],
        [1900, 3740, 3720],
        first_col_bold=True,
    )
    add_heading = doc.add_heading("A useful README structure", level=2)
    for text in [
        "Problem chosen and why it matters",
        "Proposed mechanism, with shapes and one small worked example",
        "What can fail and what assumptions are required",
        "Baseline and fairness controls",
        "Dataset and train/validation split",
        "Results table and plots",
        "Ablations and negative results",
        "Exact commands to reproduce",
        "Conclusion limited to what the evidence supports",
    ]:
        add_bullet(doc, text)
    doc.add_page_break()

    # Section 10.
    doc.add_heading("10. Common confusions, resolved", level=1)
    add_table(
        doc,
        ["Confusion", "Resolution"],
        [
            ["Token ID vs embedding", "The ID is an address. The embedding is the dense vector read or constructed from it."],
            ["Embedding dimension vs model width", "They can be equal, but need not be. A projection connects them when they differ."],
            ["Token position vs byte position", "Token position orders words/subwords in the sequence; byte position orders bytes inside one token."],
            ["Vocabulary vs Kronecker", "Kronecker is an embedding mechanism. Tokenization still decides the spans presented as tokens."],
            ["Deterministic vs contextual", "The base code is deterministic. Transformer layers still change it according to surrounding context."],
            ["Input solved means output solved", "No. V1 can remove the learned input row table while the V-way output head remains."],
            ["Cosine similarity makes reversal easy", "Only if predictions already land near well-separated valid codes and search is affordable; this is weakest early in training."],
            ["KL divergence guarantees invertibility", "It can train distributions or tolerant regions, but uniqueness and efficient decoding still need a code design."],
            ["More vocabulary is always better", "It lowers fertility but increases embedding/output cost. The optimum depends on data and hardware."],
            ["A 32-character limit", "In V1 it is effectively a 32-byte limit, which is much stricter for multi-byte scripts."],
        ],
        [2800, 6560],
        first_col_bold=True,
    )

    doc.add_heading("Final mental model", level=2)
    add_callout(doc, "Remember this", "Ordinary embedding says: “look up the learned row for this token.” Kronecker V1 says: “construct a stable code from the token's bytes and positions, then use one shared learned adapter.” V2 asks how to retain that economy while removing the fixed-window, cross-script, modality, mathematical-structure, or reverse-decoding limitation.", fill=LIGHT_TEAL, border=TEAL)

    doc.add_heading("Glossary", level=2)
    glossary = [
        ("Embedding", "A dense numeric vector used as the model's representation of an input unit."),
        ("Fertility", "How many tokens are required to represent a piece of text; lower is usually cheaper for attention."),
        ("Gather", "Selecting embedding rows by token IDs."),
        ("Kronecker product", "A structured way to combine two vectors; with one-hot byte and position vectors it uniquely identifies a position-byte pair."),
        ("Output head", "The mapping from the transformer's hidden state to one score per vocabulary token."),
        ("Rank", "The number of independent directions needed by a matrix; a low-rank factorization uses a narrow bottleneck."),
        ("Weight tying", "Reusing the input embedding weights for the output head."),
        ("Zipf distribution", "A frequency pattern with a few very common items and a long tail of rare items."),
    ]
    for term, definition in glossary:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        set_run(p.add_run(term + ". "), size=10, color=NAVY, bold=True)
        set_run(p.add_run(definition), size=10, color=INK)

    add_source_note(doc)

    # Document properties.
    props = doc.core_properties
    props.title = "From Embedding Tables to Kronecker Embeddings"
    props.subject = "Session 7 concept guide and Kronecker Embedding V2 assignment overview"
    props.author = "ERA5 Study Guide"
    props.keywords = "embeddings, Kronecker embedding, tokenization, transformer, Session 7"

    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(build_document())
